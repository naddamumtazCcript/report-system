"""DOCX export endpoint"""
from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from pathlib import Path
import logging
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

from ..config import OUTPUT_DIR

router = APIRouter()

@router.get("/export/docx/{job_id}")
async def export_docx(job_id: str):
    """Export protocol as DOCX"""
    try:
        logger.info(f"Exporting DOCX for job_id: {job_id}")
        
        # Find protocol markdown file
        protocol_file = OUTPUT_DIR / f"{job_id}_protocol.md"
        if not protocol_file.exists():
            raise HTTPException(status_code=404, detail="Protocol not found")
        
        # Read markdown
        with open(protocol_file, 'r') as f:
            content = f.read()
        
        # Create DOCX
        doc = Document()
        
        # Parse markdown and convert to DOCX elements
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line or line == '---':
                continue
            elif line.startswith('# '):
                heading = doc.add_heading(line[2:], level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('* ') or line.startswith('- '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)
        
        # Save DOCX
        docx_file = OUTPUT_DIR / f"{job_id}_protocol.docx"
        doc.save(str(docx_file))
        
        logger.info(f"DOCX generated: {docx_file}")
        
        return FileResponse(
            path=docx_file,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"protocol_{job_id}.docx"
        )
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"DOCX export failed: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"DOCX export failed: {str(e)}")
